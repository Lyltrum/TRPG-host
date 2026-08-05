"""取文：任意文件 → 纯文本 + 图片清单（`exec/29` 第 2 步）。

## 这一层以前不在代码里

`exec/26` 说模组导入「只差接线」，漏掉的正是这一整段：`scripts/module_probe/`
五个脚本**全部从 txt 开始**，而现有五份模组的 txt 全是人手工转的。实测原始文件
**95% 以上的体积是图**（死者的顿足舞 9.9 MB → 98 KB txt）。

## 判据：不是「能不能抽」，是「抽坏了能不能被发现」

四种 PDF 形态里，只有「扫描件 OCR 过、文本层有错字」那一种不可检测——对它我们
无能为力，如实说。其余三种都能给出诚实答复，而**不确定时一律显式拒绝**：
静默产出半份文本是最坏的结果，因为后面每一步都会成功，最后得到一个空壳模组。

## 🔴 图留占位，不静默丢弃

图本身不入库、不呈现（`exec/29 §5.2`：手边资料是一条独立的产品能力，本期不做），
但必须留下「这里曾有一张图」——实测四份手工 txt 里占位标记 **0 个**，丢了什么
我们完全不知道。留占位是「禁止静默兜底」的同族，也让组装时模型知道此处有图。

## 选型

`pdfplumber` + `python-docx`，两个都是 MIT。🔴 PyMuPDF 更快更全但是 **AGPL**
——本项目是网络服务，触发条件成立，且其供应商有实际诉讼记录。pdfplumber 同样
提供字级坐标（双栏按列排序可做），短板只有慢与图片导出别扭，而我们对图只清点
不存取，短板用不上。
"""

from __future__ import annotations

from dataclasses import dataclass, field
from pathlib import Path

#: 扫描件判据：每页平均可抽字符数低于此值 → 视为无文本层。
#: 🔴 判的是**密度**不是总长：200 页的扫描件光页眉页脚也能凑出几百字符，
#: 用绝对长度判会漏。这个值标定得很松——**只拦「几乎什么都没抽到」**，
#: 因为误拒一份好模组比放过一份坏模组更伤（用户手里那份就玩不了了）。
MIN_CHARS_PER_PAGE = 20

SUPPORTED_SUFFIXES = (".txt", ".pdf", ".docx", ".doc")

#: 🔴 pdfplumber 判「这两个字符之间算不算一个空格」的阈值（PDF 用户单位）。
#:
#: 默认 3 是按西文字距定的，**中文正文的字间距本来就接近 3**，于是真正的空格
#: 被一并吞掉：`爪击 70  1D6+1D6` 抽出来变成 `爪击70 1D6+1D6`。技能名与它的
#: 成功率粘成一个词，下游再也分不开——之前判成"PDF 属性表被拍平"、以为要写
#: 表格识别器，其实是这一层**没有保真搬运**。
#:
#: 2 是四份真实模组上量出来的拐点：粘连 46/34/14 → 0/0/3，同时汉字之间不会
#: 被插进多余空格（继续调到 1 就开始切碎）。剩下的少数几处是脚注上标数字，
#: 属于另一个现象，不归这个参数管。
PDF_X_TOLERANCE = 2.0


class UnsupportedDocumentError(ValueError):
    """这份文稿本层处理不了。**必须带上可执行的下一步**，不能只说"不支持"。"""


@dataclass(frozen=True, slots=True)
class ExtractedDocument:
    """取文结果。`warnings` 是**显式降级**清单，不是可以忽略的提示。"""

    text: str
    format: str
    page_count: int
    image_count: int
    warnings: list[str] = field(default_factory=list)


def image_placeholder(page: int, index: int) -> str:
    """图片占位标记。组装时模型看到它就知道此处有图。"""
    return f"[图片 第{page}页 第{index}张]"


def looks_like_scan(text: str, *, page_count: int) -> bool:
    """抽出来的东西是不是「几乎什么都没有」。"""
    pages = max(1, page_count)
    return len(text.strip()) / pages < MIN_CHARS_PER_PAGE


def extract_document(path: Path) -> ExtractedDocument:
    """按扩展名分派。不支持 / 抽不出 → 抛 `UnsupportedDocumentError`。"""
    suffix = path.suffix.lower()
    if suffix == ".txt":
        return _extract_txt(path)
    if suffix == ".docx":
        return _extract_docx(path)
    if suffix == ".pdf":
        return _extract_pdf(path)
    if suffix == ".doc":
        return _extract_legacy_doc(path)
    raise UnsupportedDocumentError(
        f"不支持的文件类型 {suffix or '（无扩展名）'}；当前支持：{'、'.join(SUPPORTED_SUFFIXES)}"
    )


def _extract_txt(path: Path) -> ExtractedDocument:
    text = path.read_text(encoding="utf-8", errors="replace")
    return ExtractedDocument(text=text, format="txt", page_count=1, image_count=0)


def _extract_docx(path: Path) -> ExtractedDocument:
    from docx import Document

    doc = Document(str(path))
    blocks = [p.text for p in doc.paragraphs]
    # 🔴 **表格必须单独读**：`doc.paragraphs` 看不见表格里的文字，而模组里
    # NPC 属性块（STR/CON/SIZ…）经常就是表格——只读段落会把它们整块丢掉。
    # 抽样的神秘渡轮表格只有 70 字所以没现形，但那是运气，不是结论。
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            if any(cells):
                blocks.append(" | ".join(cells))
    # docx 的图挂在 package part 上，按 content-type 数即可（不导出字节）
    images = sum(1 for part in doc.part.package.parts if part.content_type.startswith("image/"))
    text = "\n".join(blocks)
    warnings: list[str] = []
    if images:
        text += "\n" + "\n".join(image_placeholder(1, i + 1) for i in range(images))
        warnings.append(_handout_warning(images))
    if looks_like_scan(text, page_count=1):
        raise UnsupportedDocumentError(
            "这份 docx 里几乎没有可抽取的文字（可能整篇都是图片）。本系统不做 OCR。"
        )
    return ExtractedDocument(
        text=text, format="docx", page_count=1, image_count=images, warnings=warnings
    )


def _extract_legacy_doc(path: Path) -> ExtractedDocument:
    """Word 97-2003（OLE2）。解析细节见 `doc_legacy.py`。

    🔴 `.doc` 在中文跑团圈很常见（本仓库的科比特先生就是），**不能拒绝**。
    自己解而不调 LibreOffice/antiword，是为了不背 ~500 MB 的部署依赖；
    啃不动的情形**显式失败并给出可执行的下一步**，不返回半份文本。

    图片：`.doc` 把图埋在正文流里，本层**不清点**（清点它要再解一层结构，
    而我们对图只做告知）。这是**已知的不对称**——PDF/DOCX 会报图片数，`.doc` 不会。
    """
    from app.core.module_import.doc_legacy import LegacyDocError, extract_doc_text

    try:
        text = extract_doc_text(path)
    except LegacyDocError as exc:
        raise UnsupportedDocumentError(
            f"这份 .doc 读不出来（{exc}）。请用 Word 或 WPS 另存为 .docx 后重试。"
        ) from exc

    if looks_like_scan(text, page_count=1):
        raise UnsupportedDocumentError(
            "这份 .doc 里几乎没有可抽取的文字（可能整篇都是图片）。本系统不做 OCR。"
        )
    return ExtractedDocument(
        text=text,
        format="doc",
        page_count=1,
        image_count=0,
        warnings=["`.doc` 里的图片本层不清点，图片相关提示可能缺失。"],
    )


def _extract_pdf(path: Path) -> ExtractedDocument:
    import pdfplumber

    pages_text: list[str] = []
    images = 0
    with pdfplumber.open(str(path)) as pdf:
        page_count = len(pdf.pages)
        for page_no, page in enumerate(pdf.pages, start=1):
            n_img = len(page.images)
            images += n_img
            body = page.extract_text(x_tolerance=PDF_X_TOLERANCE) or ""
            marks = "".join("\n" + image_placeholder(page_no, k + 1) + "\n" for k in range(n_img))
            pages_text.append(body + marks)

    text = "\n".join(pages_text)
    if looks_like_scan(text, page_count=page_count):
        raise UnsupportedDocumentError(
            f"这份 PDF 没有可用的文本层（{page_count} 页只抽出 {len(text.strip())} 个字符），"
            "多半是扫描件。本系统不做 OCR，请换一份带文字的版本。"
        )
    warnings = [_handout_warning(images)] if images else []
    return ExtractedDocument(
        text=text, format="pdf", page_count=page_count, image_count=images, warnings=warnings
    )


def _handout_warning(count: int) -> str:
    """🔴 显式降级，不是静默丢弃。

    手边资料（剪报、信件、照片）是 COC 的核心装置，而系统当前**没有给玩家看图
    的能力**（`ScenarioModule` 无图片字段、前后端无呈现）。所以照实说，不假装
    什么都没发生。真风险是「某条线索只印在图上、正文没写」——那条线索会消失，
    而 `validate` 查不出来。
    """
    return (
        f"这份文稿含 {count} 张图片（可能包括玩家手边资料）。"
        "本版本无法呈现图片，守秘人会用文字描述代替；"
        "若某条线索只印在图上，它可能会缺失。"
    )
