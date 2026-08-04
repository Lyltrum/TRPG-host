"""取文层：任意文件 → 纯文本 + 图片清单（`exec/29` 第 2 步）。

## 🔴 这一层在 `exec/26` 里被整条漏掉了

「模组导入只差接线」那句话漏了两条线，这是其中一条：`scripts/module_probe/` 五个
脚本**全部从 txt 开始**，后端依赖里没有任何 PDF/docx 库。**文件 → 文本这一步以前
不在代码里，是人手工做的**（现有五份模组的 txt 全是手工转的）。

实测原始文件 **95% 以上的体积是图**（死者的顿足舞 9.9 MB → 98 KB txt，100:1）。

## 判据不是「能不能抽」，是「抽坏了能不能被发现」

| 形态 | 能抽吗 | 失败可检测吗 |
|---|---|---|
| 数字排版、文本层完好 | 能 | — |
| 双栏 / 图文混排 | 能，但阅读顺序错乱 | 可测（行长分布） |
| 纯扫描图（无文本层） | 不能 | **可测：抽出近乎空 → 拒绝** |
| 扫描件 OCR 过、文本层有错字 | 看起来能 | ❌ 不可测——最危险的一种 |

所以这里的硬约束是：**不确定就显式拒绝，绝不静默产出半份文本。**

## 🔴 图留占位，不静默丢弃

实测四份手工 txt 里图片占位标记 **0 个**——图被完全丢弃，连「这里曾有一张图」
都没留下。好消息是现有五个模组本身就是「丢掉全部图也能玩」的实证；坏消息是
**我们不知道丢了什么**。留占位是「禁止静默兜底」的同族。
"""

from __future__ import annotations

import zlib
from pathlib import Path

import pytest

from app.core.module_import.extract import (
    UnsupportedDocumentError,
    extract_document,
    looks_like_scan,
)


def _minimal_pdf(text: str, *, with_text_layer: bool = True) -> bytes:
    """手搓一份最小合法 PDF。

    不引入"能写 PDF"的第三方库——测试要验的是我们的取文与拒绝逻辑，
    为此再加一个依赖不划算。`with_text_layer=False` 造的是「扫描件」：
    页面结构齐全但没有任何文字绘制指令。
    """
    stream = f"BT /F1 12 Tf 72 720 Td ({text}) Tj ET".encode("latin-1") if with_text_layer else b""
    objs = [
        b"<< /Type /Catalog /Pages 2 0 R >>",
        b"<< /Type /Pages /Kids [3 0 R] /Count 1 >>",
        b"<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] "
        b"/Resources << /Font << /F1 5 0 R >> >> /Contents 4 0 R >>",
        b"<< /Length %d >>\nstream\n%s\nendstream" % (len(stream), stream),
        b"<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>",
    ]
    out = bytearray(b"%PDF-1.4\n")
    offsets = [0]
    for i, body in enumerate(objs, start=1):
        offsets.append(len(out))
        out += b"%d 0 obj\n" % i + body + b"\nendobj\n"
    xref_at = len(out)
    out += b"xref\n0 %d\n" % (len(objs) + 1)
    out += b"0000000000 65535 f \n"
    for off in offsets[1:]:
        out += b"%010d 00000 n \n" % off
    out += b"trailer\n<< /Size %d /Root 1 0 R >>\nstartxref\n%d\n%%%%EOF\n" % (
        len(objs) + 1,
        xref_at,
    )
    assert zlib is not None  # 仅为说明本函数不做压缩
    return bytes(out)


# ── 支持的格式 ────────────────────────────────────────


def test_txt_passes_through(tmp_path: Path) -> None:
    src = tmp_path / "m.txt"
    src.write_text("第一行\n第二行\n", encoding="utf-8")

    doc = extract_document(src)

    assert doc.format == "txt"
    assert "第一行" in doc.text and "第二行" in doc.text
    assert doc.image_count == 0


def test_docx_is_extracted(tmp_path: Path) -> None:
    from docx import Document

    src = tmp_path / "m.docx"
    d = Document()
    d.add_paragraph("宅邸的门厅积满灰尘。")
    d.add_paragraph("厨房的水槽里泡着一只搪瓷碗。")
    d.save(src)

    doc = extract_document(src)

    assert doc.format == "docx"
    assert "门厅积满灰尘" in doc.text
    assert "搪瓷碗" in doc.text


def test_docx_tables_are_read_too(tmp_path: Path) -> None:
    """🔴 `doc.paragraphs` 看不见表格里的文字。

    模组里 NPC 属性块（STR/CON/SIZ…）经常就是表格，只读段落会把它们整块丢掉
    —— 而且丢得毫无声息：字数看着正常，属性没了。
    """
    from docx import Document

    src = tmp_path / "stats.docx"
    d = Document()
    d.add_paragraph("看守梅洛迪亚斯的属性：")
    table = d.add_table(rows=2, cols=3)
    table.rows[0].cells[0].text = "STR"
    table.rows[0].cells[1].text = "CON"
    table.rows[0].cells[2].text = "SIZ"
    table.rows[1].cells[0].text = "65"
    table.rows[1].cells[1].text = "70"
    table.rows[1].cells[2].text = "55"
    d.save(src)

    doc = extract_document(src)

    assert "STR" in doc.text and "65" in doc.text
    assert "SIZ" in doc.text and "55" in doc.text


def test_pdf_with_text_layer_is_extracted(tmp_path: Path) -> None:
    src = tmp_path / "m.pdf"
    # 写实一点：真实一页文字有几百字符，13 个字符会撞上扫描件阈值（那是
    # fixture 不真实，不是阈值错）。
    body = "The Corbitt house stands at the end of a quiet lane, its windows dark."
    src.write_bytes(_minimal_pdf(body))

    doc = extract_document(src)

    assert doc.format == "pdf"
    assert "Corbitt" in doc.text
    assert doc.page_count == 1


# ── 🔴 显式拒绝：不确定就不产出 ────────────────────────


def test_broken_doc_fails_loudly_with_actionable_message(tmp_path: Path) -> None:
    """`.doc` 是**支持**的格式；但啃不动的那份要显式失败并给出下一步。

    🔴 不返回半份文本 —— 那样后面每一步都会成功，最后得到一个空壳模组。
    """
    src = tmp_path / "m.doc"
    src.write_bytes(b"\xd0\xcf\x11\xe0 not really a word document")

    with pytest.raises(UnsupportedDocumentError) as exc:
        extract_document(src)

    assert ".docx" in str(exc.value), "要给出可执行的下一步"


def test_doc_is_a_supported_suffix() -> None:
    """`.doc` 在中文跑团圈很常见（本仓库的科比特先生就是），不能拒绝。"""
    from app.core.module_import.extract import SUPPORTED_SUFFIXES

    assert ".doc" in SUPPORTED_SUFFIXES


def test_scanned_pdf_is_rejected_not_silently_empty(tmp_path: Path) -> None:
    """🔴 扫描件必须**显式失败**，不能返回一份近乎空的文本让管线继续跑。

    静默产出半份文本是最坏的结果：后面每一步都会成功，最后得到一个空壳模组。
    """
    src = tmp_path / "scan.pdf"
    src.write_bytes(_minimal_pdf("", with_text_layer=False))

    with pytest.raises(UnsupportedDocumentError) as exc:
        extract_document(src)

    assert "扫描" in str(exc.value) or "文本层" in str(exc.value)


def test_unknown_extension_is_rejected(tmp_path: Path) -> None:
    src = tmp_path / "m.epub"
    src.write_bytes(b"whatever")

    with pytest.raises(UnsupportedDocumentError):
        extract_document(src)


# ── 扫描件判据本身 ────────────────────────────────────


def test_scan_detection_is_about_density_not_absolute_length() -> None:
    """判据是「每页平均字数」，不是「总字数」。

    一份 200 页的扫描件可能因为页眉页脚也能抽出几百个字符——用绝对长度判会漏。
    """
    assert looks_like_scan("", page_count=1)
    assert looks_like_scan("页码 1 页码 2 页码 3", page_count=50)
    assert not looks_like_scan("正" * 400, page_count=1)


# ── 图片：清点 + 留占位 ───────────────────────────────


def test_image_placeholders_are_left_not_dropped(tmp_path: Path) -> None:
    """🔴 图不入库、不呈现，但必须留下「这里曾有一张图」。

    实测四份手工 txt 里占位标记 0 个——丢了什么我们完全不知道。
    """
    from app.core.module_import.extract import image_placeholder

    mark = image_placeholder(page=3, index=2)

    assert "3" in mark and "2" in mark
    assert mark.strip(), "占位不能是空串"


# ── `.doc` 分片表：最容易出错的那一段 ──────────────────


def _clx(pieces: list[tuple[int, int]], cps: list[int]) -> bytes:
    """按 Word 的格式拼一份 Clx：`0x02` + 长度 + (n+1 个 CP) + (n 个 PCD)。"""
    import struct

    body = b"".join(struct.pack("<I", cp) for cp in cps)
    for _fc, fc_field in pieces:
        body += struct.pack("<H", 0) + struct.pack("<I", fc_field) + struct.pack("<H", 0)
    return b"\x02" + struct.pack("<I", len(body)) + body


def test_piece_table_is_parsed_into_ranges() -> None:
    from app.core.module_import.doc_legacy import _parse_piece_table

    clx = _clx([(0, 0x0400), (0, 0x0800)], [0, 10, 25])

    assert _parse_piece_table(clx) == [(0, 10, 0x0400), (10, 25, 0x0800)]


def test_prc_entries_are_skipped_before_the_piece_table() -> None:
    """Clx 里可能先有若干 `0x01`（Prc）条目——跳不过去就读到垃圾。"""
    import struct

    from app.core.module_import.doc_legacy import _parse_piece_table

    prc = b"\x01" + struct.pack("<H", 4) + b"\x00\x00\x00\x00"
    clx = prc + _clx([(0, 0x0200)], [0, 7])

    assert _parse_piece_table(clx) == [(0, 7, 0x0200)]


def test_compressed_flag_decides_encoding_not_guesswork() -> None:
    """🔴 PCD.fc 第 30 位是压缩标志：置位=单字节，未置位=UTF-16LE。

    漏掉它，中文文档会抽出乱码或长度翻倍的鬼字符 —— **而且不会报错**。
    这里直接对常量断言，防止有人"顺手"改掉。
    """
    from app.core.module_import.doc_legacy import _FC_ADDRESS_MASK, _FC_COMPRESSED

    assert _FC_COMPRESSED == 1 << 30
    assert _FC_ADDRESS_MASK == (1 << 30) - 1
    assert _FC_COMPRESSED & _FC_ADDRESS_MASK == 0, "标志位不能落进地址位里"


def test_word_control_characters_become_readable_breaks() -> None:
    from app.core.module_import.doc_legacy import _clean

    assert _clean("第一段\r第二段") == "第一段\n第二段"
    assert _clean("单元格\x07") == "单元格\n"
    assert _clean("域\x13代码\x14显示\x15") == "域代码显示"
